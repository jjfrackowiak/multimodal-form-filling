"""Compile — `DerivativeArtifact` / `NetNewArtifact` → `.docx` bytes + `RenderMap`.

Comments are **not** attached here — `attach_comments` is a separate call over the bytes
either function returns. Splitting the two means the byte-identical-body promise for
derivative is a property of `compile_derivative` alone, provable with zero `ReviewComment`s
in play: it returns `source` untouched and a map of where things are, nothing more.

Net-new is "the easy one" (per the brief) precisely because there is no existing structure
to preserve — the compiler is free to build the document *and* the map in the same pass,
recording each entry's span the moment it writes it.
"""

from __future__ import annotations

import io

from docx import Document
from docx.shared import Inches

from mff_contracts import DerivativeArtifact, NetNewArtifact, RenderMap, RunSpan

from ._io import dump_document
from .parse import walk

__all__ = ["compile_derivative", "compile_netnew"]


def compile_derivative(artifact: DerivativeArtifact, source: bytes) -> tuple[bytes, RenderMap]:
    """Build the `RenderMap` for a client's document without touching it.

    The returned bytes are `source`, unchanged — a derivative compile never mutates the
    client's document; only `attach_comments` does, and only by adding comment markup.
    Re-walking `source` (rather than trusting `artifact.nodes` blindly) is what guarantees
    the ids in the returned map agree with what `parse_docx` would produce from these exact
    bytes right now, not whatever `artifact.nodes` said when it was last persisted.
    """
    # `artifact.nodes`/`artifact.comments` are not needed to build the map — see docstring.
    _, _, spans = walk(source)
    return source, RenderMap(anchor_to_span=spans)


def compile_netnew(
    artifact: NetNewArtifact,
    *,
    title: str | None = None,
    image_bytes: dict[str, bytes] | None = None,
    vehicle_fields: list[tuple[str, str]] | None = None,
    extra_images: list[tuple[str, bytes]] | None = None,
) -> tuple[bytes, RenderMap]:
    """Render a `FormDraft` from nothing, recording each entry's span as it is written.

    A `Section` has no `Anchor.kind` of its own (only `"node"`, `"entry"` and `"document"`
    exist), so section headings are written for structure but never enter the map — a
    review comment can only ever target one of its entries.

    `image_bytes` is keyed by blob sha256. When omitted, image refs stay as text
    placeholders so this module still has no I/O.
    """
    document = Document()
    spans: dict[str, RunSpan] = {}
    blobs = image_bytes or {}

    document.add_heading(title or artifact.form_id, level=0)
    if title and artifact.form_id and title != artifact.form_id:
        document.add_paragraph(artifact.form_id)

    if vehicle_fields:
        table = document.add_table(rows=len(vehicle_fields), cols=2)
        table.style = "Table Grid"
        for row, (label, value) in zip(table.rows, vehicle_fields, strict=True):
            row.cells[0].text = label
            row.cells[1].text = value

    for section in artifact.draft.sections:
        document.add_heading(section.title, level=1)
        for entry in section.entries:
            paragraph = document.add_paragraph()
            if entry.value:
                paragraph.add_run(entry.value)
            paragraph_index = len(document.paragraphs) - 1
            for image in entry.images:
                data = blobs.get(image.sha256)
                if data:
                    _add_picture(document, data)
                else:
                    paragraph.add_run(f"  [image: {image.sha256[:12]}]")
            if not paragraph.runs:
                paragraph.add_run("")
            spans[entry.id] = RunSpan(
                paragraph_index=paragraph_index,
                run_start=0,
                run_end=max(0, len(paragraph.runs) - 1),
            )

    if extra_images:
        document.add_heading("Photographs", level=1)
        for caption, data in extra_images:
            document.add_paragraph(caption)
            _add_picture(document, data)

    document.add_paragraph("Returning driver signature: ..................................")
    document.add_paragraph("Receiving officer signature: ..................................")

    return dump_document(document), RenderMap(anchor_to_span=spans)


def _add_picture(document: Document, data: bytes) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    try:
        run.add_picture(io.BytesIO(data), width=Inches(5))
    except Exception:
        run.text = "[image]"
