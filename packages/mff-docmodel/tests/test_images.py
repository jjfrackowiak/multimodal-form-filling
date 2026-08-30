"""Embedded media extraction: the common single-image-per-paragraph shape (the fixture),
the multi-image-per-paragraph edge the fixture never exercises, and the defensive path in
`_images` that keeps a non-image drawing from becoming a broken "image" node."""

from __future__ import annotations

import io

from docx import Document
from PIL import Image

from mff_docmodel._images import _is_decodable_image, _resolve, embedded_images_in_run
from mff_docmodel.parse import parse_docx


def _png_bytes(color: tuple[int, int, int]) -> bytes:
    buf = io.BytesIO()
    Image.new("RGB", (4, 4), color=color).save(buf, format="PNG")
    return buf.getvalue()


def test_is_decodable_image_true_for_real_image() -> None:
    assert _is_decodable_image(_png_bytes((10, 20, 30))) is True


def test_is_decodable_image_false_for_garbage() -> None:
    assert _is_decodable_image(b"not an image, just some bytes") is False


def test_embedded_images_in_run_empty_when_no_blip() -> None:
    document = Document()
    paragraph = document.add_paragraph("plain text, no drawing")
    run = paragraph.runs[0]
    assert embedded_images_in_run(run.element, document) == []


def test_multi_image_paragraph_gets_one_node_per_run() -> None:
    """Two different photos in two runs of the same paragraph — the edge `_image_nodes`
    exists for. Ids disambiguate with `.i{k}`; the common single-image case (the fixture)
    stays plain `p{index}`."""
    document = Document()
    paragraph = document.add_paragraph()
    run_a = paragraph.add_run()
    run_a.add_picture(io.BytesIO(_png_bytes((255, 0, 0))), width=100)
    run_b = paragraph.add_run()
    run_b.add_picture(io.BytesIO(_png_bytes((0, 255, 0))), width=100)

    buf = io.BytesIO()
    document.save(buf)

    nodes = parse_docx(buf.getvalue())
    images = [n for n in nodes if n.kind == "image"]
    assert len(images) == 2
    assert {n.id for n in images} == {"p0.i0", "p0.i1"}
    assert images[0].image_sha256 != images[1].image_sha256


def test_single_image_paragraph_gets_the_plain_id() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(_png_bytes((1, 2, 3))), width=100)

    buf = io.BytesIO()
    document.save(buf)

    nodes = parse_docx(buf.getvalue())
    images = [n for n in nodes if n.kind == "image"]
    assert len(images) == 1
    assert images[0].id == "p0"


# --- Duck-typed stand-ins for the defensive branches: a dangling relationship id and a
# --- relationship that resolves to something that is not a decodable image. Exercising
# --- these through the full python-docx API would mean hand-corrupting a real .docx; a
# --- minimal fake of the two attributes `_resolve` actually reads is far more direct. ---


class _FakePart:
    def __init__(self, blob: bytes) -> None:
        self.blob = blob


class _FakeDocumentPart:
    def __init__(self, related_parts: dict[str, _FakePart]) -> None:
        self.related_parts = related_parts


class _FakeDocument:
    def __init__(self, related_parts: dict[str, _FakePart]) -> None:
        self.part = _FakeDocumentPart(related_parts)


class _FakeBlip:
    def __init__(self, rid: str | None) -> None:
        self._rid = rid

    def get(self, _key: str) -> str | None:
        return self._rid


class _FakeRun:
    def __init__(self, blips: list[_FakeBlip]) -> None:
        self._blips = blips

    def xpath(self, _expr: str) -> list[_FakeBlip]:
        return self._blips


def test_resolve_dangling_relationship_id_returns_none() -> None:
    document = _FakeDocument(related_parts={})
    assert _resolve("rIdMissing", document) is None  # type: ignore[arg-type]


def test_resolve_non_image_part_returns_none() -> None:
    document = _FakeDocument(related_parts={"rId9": _FakePart(b"not an image")})
    assert _resolve("rId9", document) is None  # type: ignore[arg-type]


def test_embedded_images_in_run_skips_blip_without_embed_attribute() -> None:
    run = _FakeRun(blips=[_FakeBlip(rid=None)])
    document = _FakeDocument(related_parts={})
    assert embedded_images_in_run(run, document) == []  # type: ignore[arg-type]
