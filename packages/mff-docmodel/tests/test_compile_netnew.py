"""`compile_netnew`: "the easy one" — build the document and the map in the same pass."""

from __future__ import annotations

import io

from docx import Document

from mff_contracts import BlobRef, Entry, FormDraft, NetNewArtifact, Section
from mff_docmodel import compile_netnew


def _artifact() -> NetNewArtifact:
    draft = FormDraft(
        sections=[
            Section(
                id="s1",
                title="Dane pojazdu",
                entries=[
                    Entry(id="e1", order="a0", value="Nissan Qashqai", set_by="R-01"),
                    Entry(id="e2", order="a1", value="WN 7020U", set_by="R-02"),
                ],
            ),
            Section(
                id="s2",
                title="Notes",
                entries=[Entry(id="e3", order="a0", value="Stan dobry", set_by="R-03")],
            ),
        ]
    )
    return NetNewArtifact(job_id="j-1", form_id="WN-7020U", draft=draft)


def test_compile_netnew_opens_in_word() -> None:
    compiled_bytes, _render_map = compile_netnew(_artifact())
    document = Document(io.BytesIO(compiled_bytes))
    texts = [p.text for p in document.paragraphs]
    assert "WN-7020U" in texts
    assert "Dane pojazdu" in texts
    assert "Nissan Qashqai" in texts
    assert "Notes" in texts
    assert "Stan dobry" in texts


def test_render_map_has_one_span_per_entry() -> None:
    artifact = _artifact()
    _compiled_bytes, render_map = compile_netnew(artifact)
    expected_ids = {e.id for s in artifact.draft.sections for e in s.entries}
    assert set(render_map.anchor_to_span) == expected_ids


def test_render_map_spans_resolve_to_the_right_text() -> None:
    artifact = _artifact()
    compiled_bytes, render_map = compile_netnew(artifact)
    document = Document(io.BytesIO(compiled_bytes))

    span = render_map.anchor_to_span["e1"]
    paragraph = document.paragraphs[span.paragraph_index]
    run_text = "".join(r.text for r in paragraph.runs[span.run_start : span.run_end + 1])
    assert run_text == "Nissan Qashqai"


def test_entry_images_are_noted_without_needing_bytes() -> None:
    """`compile_netnew` never receives blob bytes (no network, no `BlobStore` access from
    this layer) — an entry's images are surfaced as a reference, not embedded pixels."""
    image = BlobRef(
        uri="gs://bucket/jobs/1/img/abc123",
        content_type="image/jpeg",
        size_bytes=100,
        sha256="abc123def456",
    )
    draft = FormDraft(
        sections=[
            Section(
                id="s1",
                title="Photos",
                entries=[
                    Entry(id="e1", order="a0", value="Podsufitka", images=[image], set_by="R-04")
                ],
            )
        ]
    )
    compiled_bytes, render_map = compile_netnew(
        NetNewArtifact(job_id="j-1", form_id="f", draft=draft)
    )
    document = Document(io.BytesIO(compiled_bytes))
    span = render_map.anchor_to_span["e1"]
    paragraph = document.paragraphs[span.paragraph_index]
    assert "abc123def456"[:12] in paragraph.text


def test_compile_uses_display_title_and_embeds_image_bytes() -> None:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (8, 8), (10, 20, 30)).save(buf, format="JPEG")
    jpeg = buf.getvalue()
    image = BlobRef(
        uri="gs://bucket/image/abc",
        content_type="image/jpeg",
        size_bytes=len(jpeg),
        sha256="abc123def456",
    )
    draft = FormDraft(
        sections=[
            Section(
                id="s1",
                title="1. Under the bonnet",
                entries=[
                    Entry(id="e1", order="a0", value="Engine bay", images=[image], set_by="R-01")
                ],
            )
        ]
    )
    compiled, _map = compile_netnew(
        NetNewArtifact(job_id="job-9", form_id="job-9", draft=draft),
        title="Vehicle return report",
        image_bytes={"abc123def456": jpeg},
        extra_images=[("spare.jpg", jpeg)],
    )
    document = Document(io.BytesIO(compiled))
    texts = [p.text for p in document.paragraphs]
    assert "Vehicle return report" in texts
    assert "1. Under the bonnet" in texts
    assert "Photographs" in texts
    assert document.inline_shapes


def test_empty_draft_still_produces_an_openable_document() -> None:
    compiled_bytes, render_map = compile_netnew(
        NetNewArtifact(job_id="j-1", form_id="empty", draft=FormDraft())
    )
    document = Document(io.BytesIO(compiled_bytes))
    assert document.paragraphs
    assert render_map.anchor_to_span == {}
