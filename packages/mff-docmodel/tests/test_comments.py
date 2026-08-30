"""`attach_comments` — the `RenderMap` bridge, exercised against both compilers.

python-docx >= 1.2 is required for any of this: `Document.add_comment` and
`Document.comments` do not exist at all on 1.1.0 (verified — see the package README).
"""

from __future__ import annotations

import io

from docx import Document as OpenDocument

from mff_contracts import Anchor, BlobRef, DerivativeArtifact, ReviewComment
from mff_docmodel import attach_comments, compile_derivative, parse_docx

AUTHOR = "AI Editor"


def _artifact(source: bytes) -> DerivativeArtifact:
    nodes = parse_docx(source)
    blob = BlobRef(
        uri="mem://s", content_type="application/vnd.docx", size_bytes=len(source), sha256="x"
    )
    return DerivativeArtifact(job_id="j-1", form_id="form_supplied", source=blob, nodes=nodes)


def _comment(
    requirement_id: str, target_id: str, *, verdict: str = "pass", suggestion: str | None = None
) -> ReviewComment:
    return ReviewComment(
        requirement_id=requirement_id,
        anchor=Anchor(kind="node", target_id=target_id),
        verdict=verdict,
        justification=f"justification for {requirement_id}, long enough to pass the length check.",
        suggestion=suggestion,
    )


def test_attach_single_comment(derivative_docx_bytes: bytes) -> None:
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)
    heading = next(n for n in artifact.nodes if n.text == "6. Tyre tread")

    out_bytes, count, unanchored = attach_comments(
        compiled_bytes, [_comment("R-07", heading.id)], render_map, author=AUTHOR
    )

    assert count == 1
    assert unanchored == []
    doc = OpenDocument(io.BytesIO(out_bytes))
    comments = list(doc.comments)
    assert len(comments) == 1
    assert comments[0].author == AUTHOR
    assert "[R-07]" in comments[0].text
    assert "PASS" in comments[0].text


def test_two_comments_on_one_span_the_r05_r06_shape(derivative_docx_bytes: bytes) -> None:
    """The fixture's real case: R-05 and R-06 both target the same heading node."""
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)
    heading = next(n for n in artifact.nodes if n.text == "5. Windscreen")

    comments = [
        _comment("R-05", heading.id),
        _comment("R-06", heading.id),
    ]
    out_bytes, count, unanchored = attach_comments(
        compiled_bytes, comments, render_map, author=AUTHOR
    )

    assert count == 2
    assert unanchored == []
    doc = OpenDocument(io.BytesIO(out_bytes))
    doc_comments = list(doc.comments)
    assert len(doc_comments) == 2
    texts = {c.text for c in doc_comments}
    assert any("[R-05]" in t for t in texts)
    assert any("[R-06]" in t for t in texts)

    # Both comments must anchor to the SAME paragraph (the heading), not two different ones.
    heading_paragraph = next(p for p in doc.paragraphs if p.text == "5. Windscreen")
    xml = heading_paragraph._p.xml
    assert xml.count("commentRangeStart") == 2
    assert xml.count("commentRangeEnd") == 2


def test_document_anchored_comment_lands_and_is_reported(derivative_docx_bytes: bytes) -> None:
    """DoD 5: `unverified`'s legitimate home — never dropped."""
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)

    comment = ReviewComment(
        requirement_id="R-99",
        anchor=Anchor(kind="document"),
        verdict="unverified",
        justification="Could not determine what this requirement refers to in the form.",
    )
    out_bytes, count, unanchored = attach_comments(
        compiled_bytes, [comment], render_map, author=AUTHOR
    )

    assert count == 1
    assert unanchored == ["R-99"]
    doc = OpenDocument(io.BytesIO(out_bytes))
    comments = list(doc.comments)
    assert len(comments) == 1
    assert "[R-99]" in comments[0].text
    assert "UNVERIFIED" in comments[0].text


def test_unresolvable_node_id_also_falls_back_and_is_reported(derivative_docx_bytes: bytes) -> None:
    """Defensive path: a `"node"`-kind anchor whose target id is not in the map (stale
    artifact, id drift) must not raise or silently vanish — same fallback as `"document"`."""
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)

    comment = _comment("R-42", "p9999-does-not-exist")
    out_bytes, count, unanchored = attach_comments(
        compiled_bytes, [comment], render_map, author=AUTHOR
    )

    assert count == 1
    assert unanchored == ["R-42"]
    doc = OpenDocument(io.BytesIO(out_bytes))
    assert len(list(doc.comments)) == 1


def test_mixed_batch_counts_and_unanchored_list(derivative_docx_bytes: bytes) -> None:
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)
    heading = next(n for n in artifact.nodes if n.text == "8. Gauges")

    comments = [
        _comment("R-10", heading.id),
        ReviewComment(
            requirement_id="R-11",
            anchor=Anchor(kind="document"),
            verdict="unverified",
            justification="No identifiable target for this requirement in the submitted form.",
        ),
        _comment(
            "R-12",
            heading.id,
            verdict="fail",
            suggestion="Fix it by doing the thing that is suggested here, in detail.",
        ),
    ]
    out_bytes, count, unanchored = attach_comments(
        compiled_bytes, comments, render_map, author=AUTHOR
    )

    assert count == 3
    assert unanchored == ["R-11"]
    doc = OpenDocument(io.BytesIO(out_bytes))
    assert len(list(doc.comments)) == 3


def test_fail_verdict_carries_its_suggestion(derivative_docx_bytes: bytes) -> None:
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)
    heading = next(n for n in artifact.nodes if n.text == "1. Under the bonnet")

    comment = _comment(
        "R-01", heading.id, verdict="fail", suggestion="Supply a second engine-bay photograph."
    )
    out_bytes, _count, _unanchored = attach_comments(
        compiled_bytes, [comment], render_map, author=AUTHOR
    )
    doc = OpenDocument(io.BytesIO(out_bytes))
    text = next(iter(doc.comments)).text
    assert "FAIL" in text
    assert "Suggestion:" in text
    assert "Supply a second engine-bay photograph." in text
