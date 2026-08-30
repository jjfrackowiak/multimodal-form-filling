"""`compile_derivative`: builds the `RenderMap`, never touches the client's document."""

from __future__ import annotations

import io

from docx import Document
from mff_contracts import BlobRef, DerivativeArtifact

from mff_docmodel import compile_derivative, parse_docx
from mff_docmodel.parse import walk

from _xml import canonical, read_document_xml


def _artifact(source: bytes) -> DerivativeArtifact:
    nodes = parse_docx(source)
    blob = BlobRef(
        uri="mem://source",
        content_type="application/vnd.docx",
        size_bytes=len(source),
        sha256="x",
    )
    return DerivativeArtifact(form_id="form_supplied", source=blob, nodes=nodes)


def test_compile_returns_source_bytes_unchanged(derivative_docx_bytes: bytes) -> None:
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, _render_map = compile_derivative(artifact, derivative_docx_bytes)
    assert compiled_bytes == derivative_docx_bytes


def test_compile_opens_in_word(derivative_docx_bytes: bytes) -> None:
    """"Opens in Word" stand-in: python-docx can load it back without error and the
    structure is intact."""
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, _render_map = compile_derivative(artifact, derivative_docx_bytes)
    document = Document(io.BytesIO(compiled_bytes))
    assert len(document.paragraphs) == 50
    assert len(document.tables) == 1


def test_render_map_covers_every_node_that_has_a_span(derivative_docx_bytes: bytes) -> None:
    """Table cells are the one node kind with no `RunSpan` — see `_table_cell_nodes`."""
    artifact = _artifact(derivative_docx_bytes)
    nodes = parse_docx(derivative_docx_bytes)
    _compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)

    spannable = [n for n in nodes if n.kind != "table_cell"]
    assert spannable  # sanity
    for node in spannable:
        assert node.id in render_map.anchor_to_span

    table_cells = [n for n in nodes if n.kind == "table_cell"]
    assert table_cells
    for cell in table_cells:
        assert cell.id not in render_map.anchor_to_span


def test_render_map_ids_agree_with_a_fresh_parse(derivative_docx_bytes: bytes) -> None:
    """The map `compile_derivative` builds must resolve the *same* ids `parse_docx`
    reports from these exact bytes — that is the whole point of sharing `walk()`."""
    artifact = _artifact(derivative_docx_bytes)
    nodes = parse_docx(derivative_docx_bytes)
    _compiled_bytes, render_map = compile_derivative(artifact, derivative_docx_bytes)

    spannable_ids = {n.id for n in nodes if n.kind != "table_cell"}
    assert spannable_ids == set(render_map.anchor_to_span)


def test_render_map_spans_resolve_to_real_runs(derivative_docx_bytes: bytes) -> None:
    document = Document(io.BytesIO(derivative_docx_bytes))
    _walked = walk(derivative_docx_bytes)
    for node_id, span in _walked.spans.items():
        paragraph = document.paragraphs[span.paragraph_index]
        assert span.run_start <= span.run_end < len(paragraph.runs), node_id


def test_byte_identical_body_with_zero_comments(derivative_docx_bytes: bytes) -> None:
    """DoD 3, the strong form: with no comments in play at all, the compiled body is not
    merely *comment-strippable* to the original — it IS the original, byte for byte."""
    artifact = _artifact(derivative_docx_bytes)
    compiled_bytes, _render_map = compile_derivative(artifact, derivative_docx_bytes)

    before = canonical(read_document_xml(derivative_docx_bytes))
    after = canonical(read_document_xml(compiled_bytes))
    assert before == after
