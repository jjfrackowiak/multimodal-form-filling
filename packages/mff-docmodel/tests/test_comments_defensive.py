"""Defensive branches in `attach_comments` that the fixture never needs but must still be
correct: out-of-range spans (a stale `RenderMap`) and the pathological empty document."""

from __future__ import annotations

import io

from docx import Document
from mff_contracts import Anchor, RenderMap, ReviewComment, RunSpan

from mff_docmodel import attach_comments
from mff_docmodel.comments import _fallback_run_pair, _resolve

AUTHOR = "AI Editor"


def _comment(requirement_id: str, target_id: str) -> ReviewComment:
    return ReviewComment(
        requirement_id=requirement_id,
        anchor=Anchor(kind="node", target_id=target_id),
        verdict="pass",
        justification="a long enough justification to satisfy the model validator here.",
    )


def test_resolve_out_of_range_paragraph_index_falls_back() -> None:
    document = Document()
    document.add_paragraph("only paragraph")
    render_map = RenderMap(anchor_to_span={"x": RunSpan(paragraph_index=5, run_start=0, run_end=0)})
    assert _resolve(_comment("R-1", "x"), render_map, document) is None


def test_resolve_out_of_range_run_index_falls_back() -> None:
    document = Document()
    document.add_paragraph("only paragraph")
    render_map = RenderMap(anchor_to_span={"x": RunSpan(paragraph_index=0, run_start=9, run_end=9)})
    assert _resolve(_comment("R-1", "x"), render_map, document) is None


def test_fallback_run_pair_none_for_a_truly_empty_document() -> None:
    assert _fallback_run_pair(Document()) is None


def test_attach_comments_on_empty_document_drops_nothing_it_can_place_but_does_not_crash() -> None:
    document = Document()
    buf = io.BytesIO()
    document.save(buf)
    comment = _comment("R-1", "does-not-exist")

    out_bytes, count, unanchored = attach_comments(
        buf.getvalue(), [comment], RenderMap(), author=AUTHOR
    )
    # No run anywhere to anchor to: the comment cannot be attached, so it is neither
    # counted as attached nor silently reported as anchored.
    assert count == 0
    assert unanchored == []
    assert out_bytes  # still a valid document, just no comment could land
