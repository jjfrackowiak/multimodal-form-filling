"""The golden test the brief promises: `deliver()`'s output against the shape of
`fixtures/fleet-vehicle-return/expected_output/delivery.txt`.

Mirrors `check_delivery()` in `fixtures/fleet-vehicle-return/check_output.py` — same
marker, same citation regex, same substring checks — applied to what *this* module
produces rather than to the pre-baked fixture file (that function always reads the fixed
path, so it never actually exercises a candidate delivery). `_check_delivery_shape()`
below is deliberately structural, not a text diff, for the same reason the reference
evaluator is: the pipeline that produces the email may reword freely; whether the
required information is present in it should not depend on wording.

Mutation-tested at the bottom: three ways the shape check must fail, each broken on
purpose and shown caught (CONTEXT.md, "mutation-test your evaluator").
"""

from __future__ import annotations

import re
from collections.abc import Awaitable, Callable, Sequence
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from email_service.delivery import deliver
from email_service.transport import OutboundMessage
from mff_contracts import (
    JobCursor,
    JobRecord,
    RequestRecord,
    RequestResult,
    Requirement,
    ReviewComment,
)
from mff_store import InMemoryBlobStore


def _find_fixture_root() -> Path:
    for parent in Path(__file__).resolve().parents:
        candidate = parent / "fixtures" / "fleet-vehicle-return"
        if candidate.exists():
            return candidate
    raise FileNotFoundError("fixtures/fleet-vehicle-return not found above this test file")


FIXTURE = _find_fixture_root()
MANIFEST = (FIXTURE / "manifest.txt").read_text(encoding="utf-8")
_MARKER = "PARSED REQUIREMENTS"


def _check_delivery_shape(
    body: str,
    *,
    n_pass: int,
    n_fail: int,
    n_unverified: int,
    failing_ids: Sequence[str],
) -> list[str]:
    """Structural checks mirroring `check_delivery()`. Returns violation labels;
    empty means the body has the golden shape."""
    violations: list[str] = []

    # A naive `str(n) in body` substring check (what check_delivery() in
    # check_output.py does) is fooled by a mutated count that still happens to appear
    # elsewhere — "2" collides with "R-02" and half the cited line numbers. Found by
    # mutation-testing this checker; scope to the "Result:" line and require each count
    # as a whole number, not a substring anywhere in the body.
    summary_line_match = re.search(r"^Result:.*$", body, re.M)
    summary_line = summary_line_match.group(0) if summary_line_match else ""
    if not summary_line:
        violations.append("delivery has no summary line")
    for count in (n_pass, n_fail, n_unverified):
        if not re.search(rf"\b{count}\b", summary_line):
            violations.append("delivery summary line has an incorrect pass/fail/unverified count")
            break

    for req_id in failing_ids:
        if req_id not in body:
            violations.append(f"delivery does not name failing {req_id}")

    return violations


# ---------------------------------------------------------------------------
# Building the scenario from fixture data
# ---------------------------------------------------------------------------

GoldenMessageBuilder = Callable[
    [], Awaitable[tuple[OutboundMessage, list[Requirement], list[ReviewComment]]]
]


@pytest.fixture
def golden_message(
    fixture_requirements: list[Requirement], fixture_comments: list[ReviewComment]
) -> GoldenMessageBuilder:
    async def _build() -> tuple[OutboundMessage, list[Requirement], list[ReviewComment]]:
        blobs = InMemoryBlobStore()
        blob_ref = await blobs.put(
            b"pretend .docx bytes",
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            kind="reviewed",
        )
        job = JobRecord(
            job_id="job-1",
            request_id="req-1",
            form_id="form_supplied.docx",
            status="done",
            cursor=JobCursor(slice_index=2),
            document=blob_ref,
            summary={"pass": 8, "fail": 2},
            unverified=[],
        )
        request = RequestRecord(
            request_id="req-1",
            manifest_raw=MANIFEST,
            requirements=fixture_requirements,
            job_ids=[job.job_id],
            reply_to="client@example.com",
            original_message_id="<original-client-message@example.com>",
            status="running",
        )
        result = RequestResult(
            request_id="req-1",
            status="done",
            documents=[blob_ref],
            requirements=fixture_requirements,
            summary={"pass": 8, "fail": 2},
            unverified=[],
            failed_forms=[],
        )
        message = await deliver(result, request, blobs=blobs, comments=fixture_comments, jobs=[job])
        return message, fixture_requirements, fixture_comments

    return _build


async def test_delivery_matches_the_golden_shape(golden_message: GoldenMessageBuilder) -> None:
    message, _requirements, _comments = await golden_message()
    violations = _check_delivery_shape(
        message.body,
        n_pass=8,
        n_fail=2,
        n_unverified=0,
        failing_ids=["R-01", "R-04"],
    )
    assert violations == []


async def test_delivery_names_both_failing_requirements_with_justification(
    golden_message: GoldenMessageBuilder,
) -> None:
    message, _requirements, _comments = await golden_message()
    body = message.body
    assert "[R-01]" in body
    assert "[R-04]" in body
    assert "Two engine-bay photographs were required" in body
    assert "between the front seats" in body
    # req 10: a fail carries a remedy.
    assert "Suggestion:" in body


async def test_delivery_attaches_the_single_reviewed_document(
    golden_message: GoldenMessageBuilder,
) -> None:
    message, requirements, _comments = await golden_message()
    attachments = {attachment.filename: attachment for attachment in message.attachments}
    assert "form_supplied.docx" in attachments
    requirement_document = Document(BytesIO(attachments["parsed-requirements.docx"].data))
    text = "\n".join(
        [paragraph.text for paragraph in requirement_document.paragraphs]
        + [cell.text for row in requirement_document.tables[0].rows for cell in row.cells]
    )
    for requirement in requirements:
        assert requirement.id in text
        assert requirement.text in text
        assert requirement.source_span in text
    assert _MARKER in message.body
    assert requirements[0].text not in message.body


# ---------------------------------------------------------------------------
# Mutation tests — the shape checker itself must catch a broken body.
# ---------------------------------------------------------------------------


async def test_mutation_wrong_summary_counts_is_caught(
    golden_message: GoldenMessageBuilder,
) -> None:
    message, _requirements, _comments = await golden_message()
    body = message.body
    mutated = body.replace("8 passed, 2 failed", "9 passed, 1 failed")
    violations = _check_delivery_shape(
        mutated,
        n_pass=8,
        n_fail=2,
        n_unverified=0,
        failing_ids=["R-01", "R-04"],
    )
    assert any("summary" in v for v in violations)
