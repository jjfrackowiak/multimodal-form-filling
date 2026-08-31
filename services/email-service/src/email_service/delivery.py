"""Delivery — the email service's **second role** (req 10, concern D2).

Intake (B3) answers "did we accept this?" synchronously. This module answers "what
happened?", minutes later, once a request's jobs have all settled. An earlier draft of
the plan forgot this role entirely: intake was covered and nothing delivered the
finished documents, so the planned system produced nothing a client could ever see.

Two collaborating pieces:

- `deliver()` — a **pure builder**. Given a settled `RequestResult` and the
  `RequestRecord` it belongs to, it decides what the reply says and what it carries
  (fetching attachment bytes or a signed link via `BlobStore`, which is the only I/O it
  performs). It never sends anything and never mutates its inputs.
- `DeliveryDispatcher` — owns *triggering* and *idempotency*. The runner calls
  `on_jobs_settled()` when it decides a request is done (that decision is B5's, not
  ours); `sweep()` is the safety net for a callback lost to a restart. Both funnel
  through the same idempotent send path: **record delivery on `RequestRecord` before
  acknowledging the send**, so a callback firing twice — or a callback racing the sweep
  — produces one email, not two.

The two roles (intake, delivery) share only the transport (`MailTransport`,
`OutboundMessage`) — see CONTEXT.md. Intake keys idempotency on `Message-ID`; this
module keys on `request_id`.

Out of scope here: deciding when a request is complete (B5), sending bytes (B4),
rendering documents (B1), what a review comment says (B6/B7). This module only
*narrates* what those stages already decided.
"""

from __future__ import annotations

import asyncio
import mimetypes
from collections.abc import Iterable, Mapping, Sequence
from dataclasses import dataclass, field
from io import BytesIO
from typing import TypeVar

from docx import Document

from mff_contracts import (
    BlobRef,
    BlobStore,
    JobRecord,
    JobRepository,
    RequestRecord,
    RequestRepository,
    RequestResult,
    Requirement,
    ReviewComment,
)

from .mail_html import render_delivery_html
from .transport.messages import Attachment, OutboundMessage
from .transport.protocol import MailTransport

__all__ = [
    "DEFAULT_ATTACH_THRESHOLD_BYTES",
    "SIGNED_URL_TTL_SECONDS",
    "DeliveryDispatcher",
    "aggregate_result",
    "deliver",
]

# Gmail caps attachments at 25 MB; "many corporate servers cap at 10" (brief). Attach
# below the tighter of the two named ceilings so a delivery never bounces regardless of
# which mail system is on the other end; link above it instead.
DEFAULT_ATTACH_THRESHOLD_BYTES = 10 * 1024 * 1024

# A signed link only has to survive long enough for the client to read the email and
# click it — a week is generous without leaving a link live indefinitely.
SIGNED_URL_TTL_SECONDS = 7 * 24 * 60 * 60

# Verdict vocabulary from mff_contracts.review.ReviewComment, split into the two
# families req 10/16 care about narrating: an "ok" outcome needs no remedy, a "needs
# attention" one is why suggestion is required at all.
_OK_VERDICTS = frozenset({"pass", "realised", "not_applicable"})
_ATTENTION_VERDICTS = frozenset({"fail", "shortfall"})

_VERDICT_LABELS = {
    "pass": "passed",
    "fail": "failed",
    "realised": "realised",
    "shortfall": "shortfalls",
    "not_applicable": "not applicable",
    "unverified": "unverified",
}

_SEPARATOR = "─" * 68
_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_REQUIREMENTS_FILENAME = "parsed-requirements.docx"


# ---------------------------------------------------------------------------
# Grouping documents by form / mode
#
# `RequestResult.documents` (frozen) is a bare `list[BlobRef]` — no form_id, no mode.
# `JobRecord` (frozen) carries `form_id` and `document` but no `mode` either: mode is
# only ever present on `JobRequest`, which does not survive to delivery. `jobs` is
# therefore an *optional* enrichment — when the caller (the runner, which already has
# these records in hand from its own completion check) supplies them, documents are
# labelled with the client's own `form_id` and grouped by mode inferred from the
# documented `form_id` convention (CONTEXT.md: ".docx filename, or input folder name").
# Without it, `deliver()` still produces a fully valid email; it just can't tell one
# document from another beyond a generic ordinal, and cannot group by mode. See the PR
# description for this gap.
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class _LabeledDocument:
    ref: BlobRef
    form_id: str | None
    mode_label: str  # "derivative" | "net_new" | "document" (unknown)


def _infer_mode_label(form_id: str) -> str:
    return "derivative" if form_id.lower().endswith(".docx") else "net_new"


def _label_documents(
    documents: Sequence[BlobRef], jobs: Sequence[JobRecord]
) -> list[_LabeledDocument]:
    by_uri: dict[str, JobRecord] = {
        job.document.uri: job for job in jobs if job.document is not None
    }
    labeled: list[_LabeledDocument] = []
    for ref in documents:
        job = by_uri.get(ref.uri)
        if job is None:
            labeled.append(_LabeledDocument(ref=ref, form_id=None, mode_label="document"))
        else:
            labeled.append(
                _LabeledDocument(
                    ref=ref, form_id=job.form_id, mode_label=_infer_mode_label(job.form_id)
                )
            )
    return labeled


def _attachment_filename(labeled: _LabeledDocument, *, index: int) -> str:
    if labeled.form_id and labeled.form_id.lower().endswith(".docx"):
        return labeled.form_id
    stem = labeled.form_id or f"document-{index}"
    extension = mimetypes.guess_extension(labeled.ref.content_type) or ".docx"
    return f"{stem}{extension}"


# ---------------------------------------------------------------------------
# Body rendering
# ---------------------------------------------------------------------------


def _render_summary_line(result: RequestResult) -> str:
    parts = []
    # `unverified` on RequestResult is authoritative (req 17); trust the list itself
    # over whatever count summary happens to carry for that key.
    for key, count in result.summary.items():
        if key == "unverified":
            continue
        label = _VERDICT_LABELS.get(key, key)
        parts.append(f"{count} {label}")
    parts.append(f"{len(result.unverified)} {_VERDICT_LABELS['unverified']}")
    if not parts:
        return "Result: no verification data."
    return "Result: " + ", ".join(parts) + "."


def _render_requirement_entry(requirement: Requirement) -> str:
    lines = [f"  {requirement.id}  {requirement.text}"]
    lines.append(
        f'        from the manifest, line {requirement.source_line}: "{requirement.source_span}"'
    )
    constraint = requirement.constraint
    if constraint is not None:
        lines.append(f"        constraint: {constraint.kind} = {constraint.value}")
        lines.append(
            f'        from the manifest, line {constraint.source_line}: "{constraint.source_span}"'
        )
    if requirement.ambiguity:
        lines.append(f"        NOTE: {requirement.ambiguity}")
    return "\n".join(lines)


def _render_requirement_list(requirements: Sequence[Requirement]) -> str:
    ordered = sorted(requirements, key=lambda r: (r.ordinal, r.text))
    entries = "\n\n".join(_render_requirement_entry(r) for r in ordered)
    return (
        "PARSED REQUIREMENTS\n\n"
        "Comments in the attached document refer to the numbers below.\n"
        "Each requirement includes the fragment of your manifest it was read from.\n\n" + entries
    )


def _requirements_attachment(requirements: Sequence[Requirement]) -> Attachment | None:
    if not requirements:
        return None
    document = Document()
    document.add_heading("Parsed requirements", level=0)
    document.add_paragraph(
        "Requirements extracted from the submitted manifest. The reviewed documents use these IDs."
    )
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("ID", "Requirement", "Source", "Constraint")):
        cell.text = text
    for requirement in sorted(requirements, key=lambda r: (r.ordinal, r.text)):
        constraint = requirement.constraint
        cells = table.add_row().cells
        cells[0].text = requirement.id
        cells[1].text = requirement.text
        cells[2].text = f'Line {requirement.source_line}: "{requirement.source_span}"'
        cells[3].text = (
            f"{constraint.kind} = {constraint.value}\n"
            f'Line {constraint.source_line}: "{constraint.source_span}"'
            if constraint is not None
            else ""
        )
        if requirement.ambiguity:
            document.add_paragraph(f"{requirement.id} note: {requirement.ambiguity}")
    output = BytesIO()
    document.save(output)
    return Attachment(
        filename=_REQUIREMENTS_FILENAME,
        content_type=_DOCX_CONTENT_TYPE,
        data=output.getvalue(),
    )


def _render_unverified_section(
    result: RequestResult, requirements_by_id: Mapping[str, Requirement]
) -> str | None:
    if not result.unverified:
        return None
    lines = [
        "UNVERIFIED",
        "",
        "The system tried three times and could not assess the requirements below.",
        "Treat them as unresolved, not as passed:",
        "",
    ]
    for req_id in result.unverified:
        requirement = requirements_by_id.get(req_id)
        text = f" {requirement.text}" if requirement else ""
        lines.append(f"  [{req_id}]{text}")
    return "\n".join(lines)


def _render_failed_forms_section(result: RequestResult) -> str | None:
    if result.status != "partial" or not result.failed_forms:
        return None
    lines = [
        "INCOMPLETE FORMS",
        "",
        "The forms below were not completed and are not attached to this",
        "message:",
        "",
    ]
    lines.extend(f"  - {form_id}" for form_id in result.failed_forms)
    return "\n".join(lines)


def _render_verdict_sections(
    comments: Sequence[ReviewComment], requirements_by_id: Mapping[str, Requirement]
) -> str | None:
    if not comments:
        return None
    attention = [c for c in comments if c.verdict in _ATTENTION_VERDICTS]
    ok = [c for c in comments if c.verdict in _OK_VERDICTS]
    sections: list[str] = []
    if attention:
        blocks = []
        for comment in sorted(attention, key=lambda c: c.requirement_id):
            block = [f"  [{comment.requirement_id}] {comment.justification}"]
            if comment.suggestion:
                block.append(f"        Suggestion: {comment.suggestion}")
            blocks.append("\n".join(block))
        sections.append("FAILED\n\n" + "\n\n".join(blocks))
    if ok:
        ids = ", ".join(c.requirement_id for c in sorted(ok, key=lambda c: c.requirement_id))
        sections.append(f"PASSED\n\n  {ids}")
    if not sections:
        return None
    return "\n\n".join(sections)


_MODE_HEADING = {
    "derivative": "reviewed forms (derivative)",
    "net_new": "composed documents (net-new)",
    "document": "documents",
}


_T = TypeVar("_T")


def _group_by_mode(
    items: Sequence[tuple[_LabeledDocument, _T]],
) -> dict[str, list[tuple[_LabeledDocument, _T]]]:
    """Group attached/linked documents by mode — the brief's "a request may mix
    modes... group them so a reader can tell which is which", applied to the one place
    delivery actually has a document to hand a reader (`jobs` is what makes the mode
    label possible at all; see the module docstring for what's missing without it)."""
    groups: dict[str, list[tuple[_LabeledDocument, _T]]] = {}
    for labeled, payload in items:
        groups.setdefault(labeled.mode_label, []).append((labeled, payload))
    return groups


def _render_attachments_section(
    attached: Sequence[tuple[_LabeledDocument, Attachment]],
    linked: Sequence[tuple[_LabeledDocument, str]],
    requirements_attachment: Attachment | None,
) -> str | None:
    if not attached and not linked and requirements_attachment is None:
        return None
    # Only worth a per-mode heading when more than one mode is actually present —
    # a single-mode request (the common case, and the fixture's) stays flat.
    modes_present = {labeled.mode_label for labeled, _ in (*attached, *linked)}
    multi_mode = len(modes_present) > 1

    lines: list[str] = []
    if attached:
        lines.append("ATTACHED DOCUMENTS")
        lines.append("")
        for mode_label, group in _group_by_mode(attached).items():
            if multi_mode:
                lines.append(f"  {_MODE_HEADING.get(mode_label, mode_label)}:")
            for labeled, attachment in group:
                kb = labeled.ref.size_bytes / 1024
                lines.append(f"  - {attachment.filename} ({kb:.0f} KB)")
    if linked:
        if lines:
            lines.append("")
        lines.append("DOCUMENTS AVAILABLE VIA LINK (over the attachment limit)")
        lines.append("")
        for mode_label, link_group in _group_by_mode(linked).items():
            if multi_mode:
                lines.append(f"  {_MODE_HEADING.get(mode_label, mode_label)}:")
            for labeled, url in link_group:
                name = labeled.form_id or "document"
                lines.append(f"  - {name}: {url}")
    if requirements_attachment is not None:
        if lines:
            lines.append("")
        lines.extend(
            [
                "PARSED REQUIREMENTS",
                "",
                f"  - {requirements_attachment.filename}",
            ]
        )
    return "\n".join(lines)


def _render_body(
    *,
    result: RequestResult,
    comments: Sequence[ReviewComment],
    attached: Sequence[tuple[_LabeledDocument, Attachment]],
    linked: Sequence[tuple[_LabeledDocument, str]],
    requirements_attachment: Attachment | None,
) -> str:
    requirements_by_id = {r.id: r for r in result.requirements}
    sections: list[str] = [
        "Hello,",
        "",
        "Your submission has been checked against "
        f"{len(result.requirements)} requirements read from your manifest.",
        "",
        _render_summary_line(result),
    ]

    verdicts = _render_verdict_sections(comments, requirements_by_id)
    if verdicts:
        sections += ["", verdicts]

    unverified_section = _render_unverified_section(result, requirements_by_id)
    if unverified_section:
        sections += ["", unverified_section]

    failed_forms_section = _render_failed_forms_section(result)
    if failed_forms_section:
        sections += ["", failed_forms_section]

    attachments_section = _render_attachments_section(attached, linked, requirements_attachment)
    if attachments_section:
        sections += ["", attachments_section]

    sections += [
        "",
        _SEPARATOR,
        "",
        "Full justifications are in the reviewer comments in the attached",
        "Word document (Review pane).",
        "",
        "--",
        "Form Validation — this message was generated automatically",
    ]
    return "\n".join(sections)


def _render_subject(result: RequestResult) -> str:
    labels = {"done": "complete", "partial": "partial", "failed": "failed"}
    status_label = labels.get(result.status, result.status)
    return f"Verification results — request {result.request_id} ({status_label})"


# ---------------------------------------------------------------------------
# The pure builder
# ---------------------------------------------------------------------------


async def _resolve_documents(
    labeled: Sequence[_LabeledDocument], *, blobs: BlobStore, threshold_bytes: int
) -> tuple[list[tuple[_LabeledDocument, Attachment]], list[tuple[_LabeledDocument, str]]]:
    attached: list[tuple[_LabeledDocument, Attachment]] = []
    linked: list[tuple[_LabeledDocument, str]] = []
    for index, doc in enumerate(labeled, start=1):
        if doc.ref.size_bytes <= threshold_bytes:
            data = await blobs.get(doc.ref)
            attachment = Attachment(
                filename=_attachment_filename(doc, index=index),
                content_type=doc.ref.content_type,
                data=data,
            )
            attached.append((doc, attachment))
        else:
            url = await blobs.signed_url(doc.ref, ttl_seconds=SIGNED_URL_TTL_SECONDS)
            linked.append((doc, url))
    return attached, linked


async def deliver(
    result: RequestResult,
    request: RequestRecord,
    *,
    blobs: BlobStore,
    comments: Sequence[ReviewComment] = (),
    jobs: Sequence[JobRecord] = (),
    attach_threshold_bytes: int = DEFAULT_ATTACH_THRESHOLD_BYTES,
) -> OutboundMessage:
    """Build the results email for a settled request (req 10).

    Pure aside from fetching attachment bytes / signed links through `blobs`: never
    sends, never mutates `result` or `request`. Threads against the client's *original*
    message (`request.original_message_id`), never our own confirmation, and marks
    `Auto-Submitted` so other automated systems do not reply to us.

    `comments` and `jobs` are optional enrichments the caller may already have (e.g. the
    runner, which assembled `result` from exactly these). Without `jobs`, documents are
    still attached/linked correctly, just without the client's own form label or a
    mode grouping. Without `comments`, the email still carries the pass/fail summary,
    every `unverified` requirement named, and the full parsed requirement list — it just
    omits the per-requirement justification narrative, which lives only in
    `ReviewComment` (B6/B7), not in `RequestResult`.
    """
    labeled = _label_documents(result.documents, jobs)
    attached, linked = await _resolve_documents(
        labeled, blobs=blobs, threshold_bytes=attach_threshold_bytes
    )
    requirements_attachment = _requirements_attachment(result.requirements)
    body = _render_body(
        result=result,
        comments=comments,
        attached=attached,
        linked=linked,
        requirements_attachment=requirements_attachment,
    )
    html_attached = [
        (attachment.filename, labeled.mode_label, labeled.ref.size_bytes)
        for labeled, attachment in attached
    ]
    html_linked = [
        (labeled.form_id or "document", labeled.mode_label, url) for labeled, url in linked
    ]
    return OutboundMessage(
        to=request.reply_to,
        subject=_render_subject(result),
        body=body,
        html_body=render_delivery_html(
            result=result,
            comments=comments,
            attached=html_attached,
            linked=html_linked,
            requirements_filename=(
                requirements_attachment.filename if requirements_attachment is not None else None
            ),
        ),
        attachments=[attachment for _, attachment in attached]
        + ([requirements_attachment] if requirements_attachment is not None else []),
        in_reply_to=request.original_message_id,
        references=[request.original_message_id],
        auto_submitted=True,
    )


# ---------------------------------------------------------------------------
# Triggering + idempotency
# ---------------------------------------------------------------------------


def aggregate_result(request: RequestRecord, records: Sequence[JobRecord]) -> RequestResult:
    """Assemble a `RequestResult` from settled `JobRecord`s.

    This is *not* "deciding when a request is complete" (B5's call, made once, on the
    happy path via the runner's own callback) — it is the sweep's safety net,
    reconstructing the same shape the runner would have handed `deliver()` directly,
    for the case where that callback never arrived (a restart mid-flight). Callers must
    only invoke this once every record in `records` has left `"running"`.
    """
    documents = [record.document for record in records if record.document is not None]
    failed_forms = [record.form_id for record in records if record.status == "failed"]
    unverified = sorted({req_id for record in records for req_id in record.unverified})
    passed = sum(
        record.summary.get("pass", 0) + record.summary.get("realised", 0) for record in records
    )
    failed = sum(
        record.summary.get("fail", 0) + record.summary.get("shortfall", 0) for record in records
    )
    status: str = "done"
    if failed_forms:
        status = "partial" if documents else "failed"
    return RequestResult(
        request_id=request.request_id,
        status=status,
        documents=documents,
        requirements=request.requirements,
        summary={"pass": passed, "fail": failed},
        unverified=unverified,
        failed_forms=failed_forms,
    )


@dataclass
class DeliveryDispatcher:
    """Owns triggering (callback + sweep) and idempotency on `request_id`.

    Two ways in, one send path:

    - `on_jobs_settled()` — the runner's completion callback.
    - `sweep()` — the safety net for a request whose callback never arrived (a restart
      mid-flight). `RequestRepository` (frozen) has no listing method, so *which*
      requests to check is the caller's job — typically whatever already tracks
      request ids (the runner, or a Firestore query the store adapter can support
      outside this frozen Protocol). `sweep()` only decides, per candidate id, whether
      its jobs have actually all settled, and if so delivers exactly like the callback
      path would have.

    Idempotency: an `asyncio.Lock` per `request_id` serialises concurrent attempts
    within this process (a callback racing the sweep, or firing twice), and the
    persisted `RequestRecord.status` is the record of what already happened — flipped
    to `"delivered"` *before* the send is acknowledged, so a crash between the two
    leaves "recorded but maybe not sent" rather than "sent twice", the direction the
    brief asks for.
    """

    requests: RequestRepository
    transport: MailTransport
    blobs: BlobStore
    attach_threshold_bytes: int = DEFAULT_ATTACH_THRESHOLD_BYTES
    _locks: dict[str, asyncio.Lock] = field(default_factory=dict, repr=False)

    def _lock_for(self, request_id: str) -> asyncio.Lock:
        lock = self._locks.get(request_id)
        if lock is None:
            lock = asyncio.Lock()
            self._locks[request_id] = lock
        return lock

    async def _deliver_once(
        self,
        result: RequestResult,
        request: RequestRecord,
        *,
        comments: Sequence[ReviewComment] = (),
        jobs: Sequence[JobRecord] = (),
    ) -> OutboundMessage | None:
        async with self._lock_for(request.request_id):
            current = await self.requests.get(request.request_id)
            if current is None:
                current = request
            if current.status != "running":
                return None  # already delivered (or never got this far) — no-op
            message = await deliver(
                result,
                current,
                blobs=self.blobs,
                comments=comments,
                jobs=jobs,
                attach_threshold_bytes=self.attach_threshold_bytes,
            )
            # Record delivery BEFORE acknowledging the send — see class docstring.
            await self.requests.put(current.model_copy(update={"status": "delivered"}))
            await self.transport.send(message)
            return message

    async def on_jobs_settled(
        self,
        result: RequestResult,
        request: RequestRecord,
        *,
        comments: Sequence[ReviewComment] = (),
        jobs: Sequence[JobRecord] = (),
    ) -> OutboundMessage | None:
        """The runner's completion callback. Returns `None` on a duplicate call."""
        return await self._deliver_once(result, request, comments=comments, jobs=jobs)

    async def sweep(
        self, request_ids: Iterable[str], *, jobs_repo: JobRepository
    ) -> list[OutboundMessage]:
        """Poll `request_ids` for ones whose jobs have all settled without a delivery
        callback ever arriving, and deliver them. Cheap: one `RequestRepository.get`
        and, only for requests still `"running"`, one `JobRepository.for_request`."""
        sent: list[OutboundMessage] = []
        for request_id in request_ids:
            record = await self.requests.get(request_id)
            if record is None or record.status != "running":
                continue
            records = await jobs_repo.for_request(request_id)
            if not records or any(job.status == "running" for job in records):
                continue  # genuinely still in flight — not overdue, or nothing to report
            result = aggregate_result(record, records)
            message = await self._deliver_once(result, record, jobs=records)
            if message is not None:
                sent.append(message)
        return sent
