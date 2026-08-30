#!/usr/bin/env python3
"""Evaluate a candidate output .docx against expected_output/structure.yaml.

The reference implementation of the L3/E1 evaluator. Deterministic, offline,
no API key. Exits 0 on zero violations, 1 otherwise.

    .venv-fixture/bin/python fixtures/fleet-vehicle-return/check_output.py \
        fixtures/fleet-vehicle-return/expected_output/report_reviewed.docx

This is what "structural, not LLM-as-judge" means in practice: the pipeline that
produced the document may be non-deterministic, but whether the document is
complete in every aspect is decided by reading it.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

import yaml
from docx import Document

HERE = Path(__file__).parent
SPEC = HERE / "expected_output" / "structure.yaml"
MANIFEST = (HERE / "manifest.txt").read_text(encoding="utf-8")
MANIFEST_LINES = MANIFEST.split("\n")

# 'line 4: "2x headliner"'  ->  ("4", "2x headliner")  — used on the DELIVERY, not comments
CITE_RE = re.compile(r'line\s+(\d+):\s*"([^"]+)"')
REQ_LINE_RE = re.compile(r"^\s*(R-\d{2})\s{2,}(.+)$", re.M)

REQ_RE = re.compile(r"\[(R-\d{2})\]")
VERDICT_RE = re.compile(r"\[R-\d{2}\]\s+(\w+)")
# Bound each field at the next field marker. Splitting on the opening marker
# alone swallows the rest of the comment and makes any field look non-empty —
# a one-character justification then passes a length check. Found by mutation
# testing the checker against a deliberately broken document.
FIELD_RE = {
    "justification": re.compile(r"Justification:\s*(.*?)(?=\n\s*(?:Suggestion:|Source:)|\Z)", re.S),
    "suggestion": re.compile(r"Suggestion:\s*(.*?)(?=\n\s*(?:Justification:|Source:)|\Z)", re.S),
}
MIN_JUSTIFICATION = 20


def field(text: str, name: str) -> str:
    m = FIELD_RE[name].search(text)
    return m.group(1).strip() if m else ""


class Report:
    def __init__(self) -> None:
        self.violations: list[str] = []
        self.checks = 0

    def check(self, ok: bool, label: str) -> None:
        self.checks += 1
        if not ok:
            self.violations.append(label)

    def __str__(self) -> str:
        head = f"{self.checks - len(self.violations)}/{self.checks} checks passed"
        if not self.violations:
            return f"PASS  {head}"
        body = "\n".join(f"  - {v}" for v in self.violations)
        return f"FAIL  {head}\n{body}"


def check_document(doc, spec: dict, r: Report) -> None:
    d = spec["document"]
    texts = [p.text.strip() for p in doc.paragraphs]

    r.check(any(t == d["title"] for t in texts), f"title missing: {d['title']}")

    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
    for h in d["required_headings"]:
        r.check(h in headings, f"heading missing: {h}")

    if d.get("headings_in_order"):
        found = [h for h in d["required_headings"] if h in headings]
        order = [headings.index(h) for h in found]
        r.check(order == sorted(order), "headings out of order")

    vt = d["vehicle_table"]
    r.check(len(doc.tables) >= 1, "vehicle table missing")
    if doc.tables:
        cells = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in doc.tables[0].rows}
        for label in vt["required_labels"]:
            r.check(label in cells, f"vehicle field missing: {label}")
        for label, value in vt.get("required_values", {}).items():
            r.check(cells.get(label) == value, f"vehicle field wrong: {label}={cells.get(label)!r}")

    # Distinct images by content hash — duplicate files must not inflate the count.
    parts = {p.sha1 for p in doc.part.package.image_parts}
    r.check(
        len(parts) >= d["min_distinct_images"],
        f"distinct images {len(parts)} < {d['min_distinct_images']}",
    )

    if d.get("signature_block_present"):
        r.check(any("signature" in t.lower() for t in texts), "signature block missing")


def check_review(doc, spec: dict, r: Report) -> None:
    rv = spec["review"]
    comments = list(doc.comments)

    r.check(
        len(comments) == rv["total_comments"],
        f"comment count {len(comments)} != {rv['total_comments']}",
    )

    seen: dict[str, str] = {}
    for c in comments:
        text = c.text
        m = REQ_RE.search(text)
        if not m:
            r.violations.append(f"comment does not name a requirement: {text[:40]!r}")
            r.checks += 1
            continue
        req = m.group(1)
        seen[req] = text

    inv = rv["invariants"]
    for req in rv["requirement_ids"]:
        r.check(req in seen, f"no comment for {req}")

    if inv.get("exactly_one_comment_per_requirement", True):
        r.check(len(seen) == len(comments), "duplicate comments for a requirement")

    for req, text in seen.items():
        vm = VERDICT_RE.search(text)
        verdict = vm.group(1).lower() if vm else ""
        expected = rv["expected_verdicts"].get(req)

        r.check(verdict == expected, f"{req} verdict {verdict!r} != {expected!r}")

        if inv.get("every_comment_has_justification"):
            just = field(text, "justification")
            r.check(
                len(just) >= MIN_JUSTIFICATION,
                f"{req} justification too short ({len(just)} chars): {just[:30]!r}",
            )

        has_sugg = bool(field(text, "suggestion"))
        if verdict == "fail" and inv.get("failing_comment_has_suggestion"):
            r.check(has_sugg, f"{req} fails without a suggestion")
        if verdict == "pass" and inv.get("passing_comment_has_no_suggestion"):
            r.check(not has_sugg, f"{req} passes but carries a suggestion")

        if inv.get("every_comment_names_its_requirement"):
            r.check(bool(REQ_RE.search(text)), f"{req} comment does not name its requirement")

        # Provenance moved to the delivery list. A comment repeating verbatim manifest
        # spans is the OLD contract and must fail, or the fixture silently keeps
        # asserting a rule the system no longer follows.
        if inv.get("no_verbatim_citation_in_comments"):
            stray = CITE_RE.findall(text)
            r.check(not stray, f"{req} carries a verbatim citation block: {stray[:1]}")

    counts: dict[str, int] = {}
    for text in seen.values():
        vm = VERDICT_RE.search(text)
        if vm:
            counts[vm.group(1).lower()] = counts.get(vm.group(1).lower(), 0) + 1
    for verdict, n in rv["verdict_counts"].items():
        r.check(counts.get(verdict, 0) == n, f"{verdict} count {counts.get(verdict, 0)} != {n}")

    unver = counts.get("unverified", 0)
    r.check(unver <= rv["unverified_allowed"], f"unverified {unver} > {rv['unverified_allowed']}")


def check_anchoring(doc, spec: dict, r: Report) -> None:
    """Comments must attach to the section they judge, not sit at the document end."""
    if not spec["review"]["invariants"].get("comment_anchored_inline"):
        return
    anchored = 0
    for p in doc.paragraphs:
        if "commentRangeStart" in p._p.xml:
            anchored += len(re.findall(r"commentRangeStart", p._p.xml))
    r.check(
        anchored == spec["review"]["total_comments"],
        f"{anchored} inline anchors for {spec['review']['total_comments']} comments",
    )


def check_ordinals(spec: dict, r: Report) -> None:
    """The ordering rule is itself an assertion: offset of source_span in the manifest."""
    expected = spec.get("expected_ordinals", {})
    if not expected:
        return
    reqs = yaml.safe_load((HERE / "expected_requirements.yaml").read_text(encoding="utf-8"))
    by_id = {x["id"]: x for x in reqs["requirements"]}
    for rid, want in expected.items():
        got = by_id.get(rid, {}).get("ordinal")
        r.check(got == want, f"{rid} ordinal {got} != {want}")
        span = by_id.get(rid, {}).get("source_span", "")
        r.check(
            span and MANIFEST.index(span) == want,
            f"{rid} ordinal is not the offset of its source_span",
        )


def check_delivery(spec: dict, r: Report) -> None:
    """Comments cite numbers; the delivery is where those numbers are explained."""
    d = spec.get("delivery")
    if not d:
        return
    path = HERE.parent.parent / d["file"] if not (HERE / d["file"]).exists() else HERE / d["file"]
    if not path.exists():
        r.check(False, f"delivery fixture missing: {d['file']}")
        return
    body = path.read_text(encoding="utf-8")

    if d.get("lists_every_requirement"):
        # Scope to the requirement-list section. Checking "rid in body" is too weak:
        # every id also appears in the pass/fail summary, so dropping one from the
        # list still passed. Found by mutation testing.
        marker = "PARSED REQUIREMENTS"
        r.check(marker in body, "delivery has no requirement-list section")
        section = body.split(marker, 1)[-1] if marker in body else ""
        listed = {m.group(1) for m in REQ_LINE_RE.finditer(section)}
        for rid in spec["review"]["requirement_ids"]:
            r.check(rid in listed, f"delivery requirement list is missing {rid}")
            # and it must carry the span it was read from, not just the number
            entry = re.search(
                rf"{rid}\s{{2,}}.+?\n(.*?)(?=\n\s*R-\d{{2}}\s{{2,}}|\Z)", section, re.S
            )
            r.check(
                bool(entry and CITE_RE.search(entry.group(1))),
                f"delivery lists {rid} without the manifest span it came from",
            )

    cited = CITE_RE.findall(body)
    r.check(bool(cited), "delivery quotes no manifest spans")
    for line_no, quote in cited:
        if d.get("quotes_are_verbatim"):
            r.check(quote in MANIFEST, f"delivery quote is not verbatim: {quote[:40]!r}")
        if d.get("cited_line_numbers_correct"):
            idx = int(line_no) - 1
            ok = 0 <= idx < len(MANIFEST_LINES) and quote in MANIFEST_LINES[idx]
            r.check(ok, f"delivery cites line {line_no} but {quote[:30]!r} is not there")

    if d.get("states_the_summary"):
        n_pass = spec["review"]["verdict_counts"]["pass"]
        n_fail = spec["review"]["verdict_counts"]["fail"]
        r.check(
            str(n_pass) in body and str(n_fail) in body,
            "delivery does not state the pass/fail counts",
        )
    if d.get("names_failing_requirements"):
        for rid, v in spec["review"]["expected_verdicts"].items():
            if v == "fail":
                r.check(rid in body, f"delivery does not name failing {rid}")


def check_coverage(spec: dict, r: Report) -> None:
    """R-04 is the case that separates a checker from a counter."""
    for c in spec["constraint_checks"]:
        r.check(
            bool(c["violated_by"]) or bool(c["satisfied_by"]),
            f"{c['requirement_id']} constraint check is empty",
        )


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        print(__doc__)
        return 2
    spec = yaml.safe_load(SPEC.read_text())["common"]
    doc = Document(argv[1])

    r = Report()
    check_document(doc, spec, r)
    check_review(doc, spec, r)
    check_anchoring(doc, spec, r)
    check_ordinals(spec, r)
    check_delivery(spec, r)
    check_coverage(spec, r)

    print(r)
    return 0 if not r.violations else 1


if __name__ == "__main__":
    sys.exit(main(sys.argv))
